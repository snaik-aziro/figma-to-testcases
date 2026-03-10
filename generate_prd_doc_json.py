import json
import sys
from pathlib import Path

def extract_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        content = {
            "source": "python-docx",
            "paragraphs": [],
            "tables": [],
            "full_text": ""
        }
        
        # Extract all text
        all_text = []
        
        # Extract paragraphs with formatting info
        for para in doc.paragraphs:
            if para.text.strip():
                content["paragraphs"].append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal"
                })
                all_text.append(para.text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content["tables"].append({
                "index": table_idx,
                "rows": table_data
            })
            # Add table content to full text
            for row in table_data:
                all_text.append(" | ".join(row))
        
        content["full_text"] = "\n".join(all_text)
        content["stats"] = {
            "total_paragraphs": len(content["paragraphs"]),
            "total_tables": len(content["tables"]),
            "character_count": len(content["full_text"])
        }
        
        return content
    except Exception as e:
        return {"error": str(e), "source": "python-docx", "type": "error"}


def extract_text_directly(file_path):
    """Simple text extraction for DOCX (reading as zip)"""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        
        content = {
            "source": "direct_xml_parsing",
            "text_blocks": [],
            "full_text": ""
        }
        
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            xml_content = zip_ref.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            # Namespace for Office Open XML
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Extract all text
            full_text = []
            for para in root.findall('.//w:p', ns):
                para_text = []
                for text_elem in para.findall('.//w:t', ns):
                    if text_elem.text:
                        para_text.append(text_elem.text)
                if para_text:
                    para_str = "".join(para_text)
                    content["text_blocks"].append(para_str)
                    full_text.append(para_str)
            
            content["full_text"] = "\n".join(full_text)
        
        return content
    except Exception as e:
        return {"error": str(e), "source": "direct_xml_parsing", "type": "error"}


def main():
    file_path = sys.argv[1] if len(sys.argv) > 1 else input("Enter file path: ").strip()
    
    if not Path(file_path).exists():
        print(f"Error: File not found: {file_path}")
        return
    
    file_ext = Path(file_path).suffix.lower()
    print(f"Extracting from {file_ext} file: {file_path}\n")
    
    results = {}
    
    if file_ext == '.docx':
        print("Method 1: Using python-docx...")
        results['docx_method'] = extract_from_docx(file_path)
        
        print("Method 2: Direct XML parsing...")
        results['xml_method'] = extract_text_directly(file_path)
    
    # Save results
    output_file = Path(file_path).stem + "_extracted.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    print(f"\nExtraction complete! Results saved to: {output_file}")
    
    # Print summary
    print("\n=== Extraction Summary ===")
    for method, data in results.items():
        if "error" not in data:
            if "stats" in data:
                stats = data["stats"]
                print(f"{method}: {stats['total_paragraphs']} paragraphs, {stats['total_tables']} tables, {stats['character_count']} characters")
            elif "full_text" in data:
                char_count = len(data["full_text"])
                block_count = len(data["text_blocks"])
                print(f"{method}: {block_count} text blocks, {char_count} characters")
        else:
            print(f"{method}: ERROR - {data['error']}")
    
    # Display first 500 characters of content
    print("\n=== Content Preview ===")
    for method, data in results.items():
        if "error" not in data and "full_text" in data:
            preview = data["full_text"][:500]
            print(f"\n{method}:\n{preview}...\n")


if __name__ == "__main__":
    main()
